#!/usr/bin/env python3
"""
generate_ci_audit_docs.py
=========================
Generates Word (.docx) documents for the CI/CD audit reports in audit-report/:

  1. audit-report/LAR_CI_Audit_Overview_2026-03-18.md
     → reports/LAR_CI_Audit_Overview_2026-03-18.docx

  2. audit-report/vertical-error-report-23225145549.txt
     → reports/LAR_Vertical_Error_Report_23225145549.docx

These companion documents are referenced in the main LAR Audit Report v7.8.
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR    = Path(__file__).parent
AUDIT_RPT   = BASE_DIR / "audit-report"
REPORTS_DIR = BASE_DIR / "reports"

CI_OVERVIEW_MD   = AUDIT_RPT / "LAR_CI_Audit_Overview_2026-03-18.md"
VERTICAL_ERR_TXT = AUDIT_RPT / "vertical-error-report-23225145549.txt"

CI_OVERVIEW_DOCX   = REPORTS_DIR / "LAR_CI_Audit_Overview_2026-03-18.docx"
VERTICAL_ERR_DOCX  = REPORTS_DIR / "LAR_Vertical_Error_Report_23225145549.docx"

# ---------------------------------------------------------------------------
# Colour palette (matches main report)
# ---------------------------------------------------------------------------
COLOUR_HEADING_DARK = RGBColor(0x1A, 0x1A, 0x2E)
COLOUR_HEADING_MID  = RGBColor(0x16, 0x21, 0x3E)
COLOUR_ACCENT       = RGBColor(0x0F, 0x3C, 0x78)
COLOUR_RED          = RGBColor(0xC0, 0x39, 0x2B)
COLOUR_ORANGE       = RGBColor(0xD3, 0x54, 0x00)
COLOUR_GREEN        = RGBColor(0x1E, 0x8B, 0x4C)
COLOUR_TABLE_HEADER = RGBColor(0x1A, 0x1A, 0x2E)


# ---------------------------------------------------------------------------
# XML / cell helpers
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_colour: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def set_paragraph_border_bottom(paragraph, colour: str = "CCCCCC"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------

def create_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)
    return doc


def add_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOUR_HEADING_DARK
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after  = Pt(6)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOUR_ACCENT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLOUR_HEADING_MID
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)

    h4 = styles["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = COLOUR_HEADING_MID
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after  = Pt(2)

    try:
        lb = styles["List Bullet"]
        lb.font.name = "Calibri"
        lb.font.size = Pt(11)
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Cover page helper
# ---------------------------------------------------------------------------

def add_cover_page(doc: Document, title: str, subtitle: str, meta: list):
    for _ in range(4):
        doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_p.add_run("OrkinosAI")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = COLOUR_ACCENT
    run.font.bold = True

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title)
    tr.font.name = "Calibri"
    tr.font.size = Pt(24)
    tr.font.bold = True
    tr.font.color.rgb = COLOUR_HEADING_DARK

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run(subtitle)
    sr.font.name = "Calibri"
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOUR_ACCENT

    doc.add_paragraph()

    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        cell_l, cell_r = row.cells[0], row.cells[1]
        set_cell_background(cell_l, "1A1A2E")
        cell_l.paragraphs[0].clear()
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(11)
        cell_r.paragraphs[0].text = value
        cell_r.paragraphs[0].runs[0].font.name = "Calibri"
        cell_r.paragraphs[0].runs[0].font.size = Pt(11)

    for row in table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Inline Markdown formatting
# ---------------------------------------------------------------------------

def _is_emoji(char: str) -> bool:
    cp = ord(char)
    return (
        (0x1F600 <= cp <= 0x1F64F)
        or (0x1F300 <= cp <= 0x1F5FF)
        or (0x1F680 <= cp <= 0x1F9FF)
        or (0x2600  <= cp <= 0x27BF)
        or (0x1FA00 <= cp <= 0x1FA9F)
        or cp == 0x200D
        or cp == 0xFE0F
    )


def clean(text: str) -> str:
    text = "".join(c for c in text if not _is_emoji(c))
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _apply_inline_formatting(run_container, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = run_container.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = run_container.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = run_container.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            if part:
                run_container.add_run(part)


# ---------------------------------------------------------------------------
# Markdown table → DOCX
# ---------------------------------------------------------------------------

def add_markdown_table(doc: Document, table_lines: list):
    rows = [line for line in table_lines
            if not re.match(r"^\|[-| :]+\|$", line.strip())]
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    n_cols = max(len(r) for r in parsed)
    n_rows = len(parsed)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row in enumerate(parsed):
        tr = table.rows[r_idx]
        for c_idx in range(n_cols):
            cell_text = clean(row[c_idx]) if c_idx < len(row) else ""
            cell = tr.cells[c_idx]
            cell.paragraphs[0].clear()
            if r_idx == 0:
                set_cell_background(cell, "1A1A2E")
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            else:
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F0F4F8")
                p = cell.paragraphs[0]
                _apply_inline_formatting(p, cell_text)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Markdown → DOCX
# ---------------------------------------------------------------------------

def md_to_docx(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        if table_buffer and not line.startswith("|"):
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        if line.startswith("|"):
            table_buffer.append(line)
            i += 1
            continue

        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            set_paragraph_border_bottom(p, "AAAAAA")
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading_text = clean(m.group(2))
            heading_text = re.sub(r"\s+#+\s*$", "", heading_text)
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        if line.startswith(">"):
            bq_text = clean(line.lstrip("> "))
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(bq_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = "Calibri"
            i += 1
            continue

        m_ol = re.match(r"^\d+\.\s+(.*)", line)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, clean(m_ol.group(1)))
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            i += 1
            continue

        m_ul = re.match(r"^[-*+]\s+(.*)", line)
        if m_ul:
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, clean(m_ul.group(1)))
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            i += 1
            continue

        m_sub = re.match(r"^\s{2,4}[-*+]\s+(.*)", line)
        if m_sub:
            p = doc.add_paragraph(style="List Bullet 2")
            _apply_inline_formatting(p, clean(m_sub.group(1)))
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            i += 1
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(cl)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(6)
        _apply_inline_formatting(p, clean(line))
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        i += 1

    if table_buffer:
        add_markdown_table(doc, table_buffer)


# ---------------------------------------------------------------------------
# Plain-text (vertical error report) → DOCX
# ---------------------------------------------------------------------------

def txt_to_docx(doc: Document, txt: str):
    """
    Convert the structured plain-text vertical error report to DOCX.
    Handles box-drawing header blocks, section dividers, and code lines.
    """
    lines = txt.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Box-drawing header blocks: ╔...╗ / ║...║ / ╚...╝
        if stripped.startswith("╔") or stripped.startswith("╚"):
            i += 1
            continue

        if stripped.startswith("║"):
            # Extract banner text from inside the box
            content = stripped.lstrip("║").rstrip("║").strip()
            if content:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(content)
                run.font.name = "Calibri"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = COLOUR_HEADING_DARK
            i += 1
            continue

        # Section dividers (═══ or ──────)
        if re.match(r"^[═─=\-]{10,}$", stripped):
            p = doc.add_paragraph()
            set_paragraph_border_bottom(p, "AAAAAA")
            i += 1
            continue

        # === Section Header === pattern
        m = re.match(r"^===\s+(.+?)\s+===$", stripped)
        if m:
            doc.add_heading(clean(m.group(1)), level=2)
            i += 1
            continue

        # --- sub-header ---
        m2 = re.match(r"^---\s+(.+?)\s+---$", stripped)
        if m2:
            doc.add_heading(clean(m2.group(1)), level=3)
            i += 1
            continue

        # Key: Value lines (e.g. "Run ID : 23225145549")
        m3 = re.match(r"^([A-Za-z ]+?)\s*[:]\s+(.+)$", stripped)
        if m3 and len(m3.group(1)) <= 25 and not stripped.startswith("PHP"):
            p = doc.add_paragraph(style="Normal")
            run_key = p.add_run(m3.group(1) + ": ")
            run_key.bold = True
            run_key.font.name = "Calibri"
            run_key.font.size = Pt(11)
            run_val = p.add_run(m3.group(2))
            run_val.font.name = "Calibri"
            run_val.font.size = Pt(11)
            i += 1
            continue

        # "--- Summary ---" style
        if stripped in ("--- Summary ---", "--- Summary---"):
            doc.add_heading("Summary", level=3)
            i += 1
            continue

        # Horizontal rule-like separator lines
        if re.match(r"^-{20,}$", stripped) or re.match(r"^={20,}$", stripped):
            p = doc.add_paragraph()
            set_paragraph_border_bottom(p, "CCCCCC")
            i += 1
            continue

        # NOTE: lines
        if stripped.upper().startswith("NOTE:"):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(stripped)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            i += 1
            continue

        # SYNTAX ERROR / WARNING lines
        if stripped.upper().startswith("SYNTAX ERROR:") or \
           stripped.upper().startswith("WARNING:"):
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOUR_RED
            i += 1
            continue

        # PHP Parse/Fatal error lines
        if "PHP Parse error" in stripped or "PHP Fatal error" in stripped or \
           "PHP Deprecated" in stripped:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOUR_RED
            i += 1
            continue

        # "Errors parsing ..." / "No syntax errors ..." lines
        if stripped.startswith("Errors parsing") or \
           stripped.startswith("No syntax errors detected"):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if stripped.startswith("Errors"):
                run.font.color.rgb = COLOUR_RED
            else:
                run.font.color.rgb = COLOUR_GREEN
            i += 1
            continue

        # "Files checked / Syntax errors / ..." lines
        if re.match(r"^(Files checked|Syntax errors|Warnings|Critical files|Missing|Present)\s*:", stripped):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            parts = stripped.split(":", 1)
            run_k = p.add_run(parts[0] + ": ")
            run_k.bold = True
            run_k.font.name = "Calibri"
            run_k.font.size = Pt(10)
            if len(parts) > 1:
                run_v = p.add_run(parts[1].strip())
                run_v.font.name = "Calibri"
                run_v.font.size = Pt(10)
            i += 1
            continue

        # TODO/FIXME lines
        if "TODO" in stripped or "FIXME" in stripped or "HACK" in stripped:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOUR_ORANGE
            i += 1
            continue

        # File path lines (start with vertical name or path separator)
        if re.match(r"^[a-z_]+/", stripped):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Default: normal paragraph
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(stripped)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        i += 1


# ---------------------------------------------------------------------------
# Build Document 1: CI Audit Overview (Markdown)
# ---------------------------------------------------------------------------

def build_ci_overview():
    print("Building: LAR CI Audit Overview 2026-03-18…")
    doc = create_document()
    add_styles(doc)

    add_cover_page(
        doc,
        title="LAR System — CI/CD Audit Overview",
        subtitle="Automated Code Quality & Deployment Audit Report",
        meta=[
            ("Report Date",           "2026-03-18"),
            ("Workflow Run ID",        "23225145549"),
            ("Commit SHA",            "185d0bb1572..."),
            ("Branch",                "main"),
            ("Prepared by",           "GitHub Actions CI/CD Pipeline"),
            ("Classification",        "CONFIDENTIAL — FOR LAR REVIEW"),
            ("Companion Document",    "LAR_Audit_Report_v7.8.docx"),
            ("Source File",           "audit-report/LAR_CI_Audit_Overview_2026-03-18.md"),
        ]
    )

    md_text = CI_OVERVIEW_MD.read_text(encoding="utf-8")
    md_to_docx(doc, md_text)

    REPORTS_DIR.mkdir(exist_ok=True)
    doc.save(CI_OVERVIEW_DOCX)
    size_kb = CI_OVERVIEW_DOCX.stat().st_size // 1024
    print(f"  Saved: {CI_OVERVIEW_DOCX}  ({size_kb} KB)")


# ---------------------------------------------------------------------------
# Build Document 2: Vertical Error Report (plain text)
# ---------------------------------------------------------------------------

def build_vertical_error_report():
    print("Building: LAR Vertical Error Report 23225145549…")
    doc = create_document()
    add_styles(doc)

    add_cover_page(
        doc,
        title="LAR System — Vertical Error Report",
        subtitle="Aggregated PHP Syntax, Config & Code Quality Findings",
        meta=[
            ("Report Date",      "2026-03-18"),
            ("Workflow Run ID",  "23225145549"),
            ("Commit SHA",       "185d0bb1572..."),
            ("Branch",           "main"),
            ("Prepared by",      "GitHub Actions CI/CD Pipeline"),
            ("Classification",   "CONFIDENTIAL — FOR LAR REVIEW"),
            ("Companion Document","LAR_Audit_Report_v7.8.docx"),
            ("Source File",      "audit-report/vertical-error-report-23225145549.txt"),
        ]
    )

    txt = VERTICAL_ERR_TXT.read_text(encoding="utf-8")
    txt_to_docx(doc, txt)

    REPORTS_DIR.mkdir(exist_ok=True)
    doc.save(VERTICAL_ERR_DOCX)
    size_kb = VERTICAL_ERR_DOCX.stat().st_size // 1024
    print(f"  Saved: {VERTICAL_ERR_DOCX}  ({size_kb} KB)")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    build_ci_overview()
    build_vertical_error_report()
    print("\nDone. Both CI audit Word documents written to reports/")
